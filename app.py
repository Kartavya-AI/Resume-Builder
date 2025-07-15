__import__('pysqlite3')
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
sys.modules["sqlite3.dbapi2"] = sys.modules["pysqlite3.dbapi2"]

import streamlit as st
import os
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from resume_builder.crew import ResumeBuilderCrew

def create_pdf_resume(resume_content):
    """Create PDF resume from content"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Split content into sections
    sections = resume_content.split('\n\n')
    
    for section in sections:
        if section.strip():
            if section.strip().isupper() or section.strip().endswith(':'):
                # Header style
                story.append(Paragraph(section, styles['Heading2']))
            else:
                # Normal text
                story.append(Paragraph(section, styles['Normal']))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def create_word_resume(resume_content):
    """Create Word resume from content"""
    doc = Document()
    
    # Split content into sections
    sections = resume_content.split('\n\n')
    
    for section in sections:
        if section.strip():
            if section.strip().isupper() or section.strip().endswith(':'):
                # Header
                heading = doc.add_heading(section, level=2)
            else:
                # Normal text
                paragraph = doc.add_paragraph(section)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(page_title="Resume Builder", page_icon="📄", layout="wide")
    
    st.title("🚀 AI Resume Builder")
    st.markdown("Generate professional resumes using AI")
    
    # Initialize session state
    if 'generate_resume' not in st.session_state:
        st.session_state.generate_resume = False
    if 'user_info' not in st.session_state:
        st.session_state.user_info = ""
    
    # Sidebar for API key
    with st.sidebar:
        st.header("⚙️ Configuration")
        api_key = st.text_input("OpenAI API Key", type="password", help="Enter your OpenAI API key")
        
        if not api_key:
            st.warning("Please enter your OpenAI API key to continue")
            st.stop()
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📝 Input Information")
        
        # Option selection
        input_option = st.radio("Choose input method:", ["Manual Entry", "Upload Raw Resume"])
        
        if input_option == "Manual Entry":
            with st.form("manual_form"):
                st.subheader("Personal Information")
                full_name = st.text_input("Full Name*")
                email = st.text_input("Email*")
                phone = st.text_input("Phone Number*")
                address = st.text_area("Address")
                
                st.subheader("Professional Information")
                job_title = st.text_input("Target Job Title*")
                experience = st.text_area("Work Experience (Previous roles, companies, duration)", height=150)
                education = st.text_area("Education (Degree, Institution, Year)", height=100)
                skills = st.text_area("Skills (Comma separated)", height=100)
                achievements = st.text_area("Key Achievements/Projects", height=100)
                
                submitted = st.form_submit_button("Generate Resume")
                
                if submitted:
                    if not all([full_name, email, phone, job_title]):
                        st.error("Please fill in all required fields marked with *")
                    else:
                        # Combine all information
                        user_info = f"""
                        Name: {full_name}
                        Email: {email}
                        Phone: {phone}
                        Address: {address}
                        Target Job Title: {job_title}
                        
                        Work Experience:
                        {experience}
                        
                        Education:
                        {education}
                        
                        Skills:
                        {skills}
                        
                        Achievements:
                        {achievements}
                        """
                        
                        st.session_state.user_info = user_info
                        st.session_state.job_title = job_title
                        st.session_state.generate_resume = True
        
        else:  # Upload Raw Resume
            uploaded_file = st.file_uploader("Upload your raw resume", type=['txt', 'doc', 'docx'])
            
            if uploaded_file is not None:
                if uploaded_file.type == "text/plain":
                    user_info = str(uploaded_file.read(), "utf-8")
                else:
                    st.error("Please upload a .txt file for now")
                    st.stop()
                
                job_title = st.text_input("Target Job Title*")
                
                if st.button("Generate Resume"):
                    if not job_title:
                        st.error("Please enter the target job title")
                    else:
                        enhanced_info = f"Target Job Title: {job_title}\n\nRaw Resume Content:\n{user_info}"
                        st.session_state.user_info = enhanced_info
                        st.session_state.job_title = job_title
                        st.session_state.generate_resume = True
    
    with col2:
        st.header("📄 Generated Resume")
        
        if st.session_state.get('generate_resume', False):
            with st.spinner("Generating your professional resume..."):
                try:
                    # Initialize CrewAI
                    crew = ResumeBuilderCrew(api_key)
                    
                    # Generate resume
                    job_title = st.session_state.get('job_title', 'Professional')
                    result = crew.run(st.session_state.user_info, job_title)
                    
                    # Display result
                    st.success("Resume generated successfully!")
                    st.text_area("Generated Resume", value=result, height=400)
                    
                    # Download options
                    st.subheader("📥 Download Options")
                    
                    col_pdf, col_word = st.columns(2)
                    
                    with col_pdf:
                        try:
                            pdf_buffer = create_pdf_resume(result)
                            st.download_button(
                                label="Download PDF",
                                data=pdf_buffer,
                                file_name="resume.pdf",
                                mime="application/pdf"
                            )
                        except Exception as e:
                            st.error(f"Error creating PDF: {str(e)}")
                    
                    with col_word:
                        try:
                            word_buffer = create_word_resume(result)
                            st.download_button(
                                label="Download Word",
                                data=word_buffer,
                                file_name="resume.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e:
                            st.error(f"Error creating Word document: {str(e)}")
                    
                    # Reset the flag
                    st.session_state.generate_resume = False
                    
                except Exception as e:
                    st.error(f"Error generating resume: {str(e)}")
                    st.error("Please check your API key and try again")
        else:
            st.info("Fill in the information on the left and click 'Generate Resume' to get started!")

if __name__ == "__main__":
    main()